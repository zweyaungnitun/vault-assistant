"""Text extraction for supported document types.

Each extractor returns the full text plus, for paginated formats (PDF), a list
of ``(page_number, char_start, char_end)`` ranges so chunks can be mapped back
to pages. Failures raise ExtractionError; ingestion logs and skips the file
rather than dropping it silently.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}

PAGE_SEPARATOR = "\n\n"


class ExtractionError(Exception):
    pass


@dataclass
class ExtractedDoc:
    text: str
    # (page_number, char_start, char_end); empty for formats without pages
    pages: list[tuple[int, int, int]] = field(default_factory=list)


def extract(path: Path) -> ExtractedDoc:
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ExtractionError(f"unsupported file type: {ext}")
    try:
        if ext in (".md", ".txt"):
            return _extract_plain(path)
        if ext == ".pdf":
            return _extract_pdf(path)
        return _extract_docx(path)
    except ExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001 — any parser failure becomes a skip+log
        raise ExtractionError(f"failed to extract {path.name}: {exc}") from exc


def _extract_plain(path: Path) -> ExtractedDoc:
    return ExtractedDoc(text=path.read_text(encoding="utf-8", errors="replace"))


def _extract_pdf(path: Path) -> ExtractedDoc:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ExtractionError(f"{path.name} is encrypted")
    parts: list[str] = []
    pages: list[tuple[int, int, int]] = []
    offset = 0
    for i, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if not page_text:
            continue
        start = offset
        parts.append(page_text)
        offset += len(page_text)
        pages.append((i, start, offset))
        parts.append(PAGE_SEPARATOR)
        offset += len(PAGE_SEPARATOR)
    text = "".join(parts).rstrip()
    if not text:
        raise ExtractionError(f"{path.name}: no extractable text (scanned/image PDF?)")
    return ExtractedDoc(text=text, pages=pages)


def _extract_docx(path: Path) -> ExtractedDoc:
    import docx

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    if not paragraphs:
        raise ExtractionError(f"{path.name}: no extractable text")
    return ExtractedDoc(text="\n\n".join(paragraphs))
