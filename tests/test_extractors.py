import pytest

from vault_assistant.extractors import ExtractionError, extract


def test_txt(tmp_path):
    f = tmp_path / "note.txt"
    f.write_text("hello world")
    doc = extract(f)
    assert doc.text == "hello world"
    assert doc.pages == []


def test_md(tmp_path):
    f = tmp_path / "note.md"
    f.write_text("# Title\n\nBody text.")
    assert "Body text." in extract(f).text


def test_docx(tmp_path):
    import docx

    f = tmp_path / "doc.docx"
    d = docx.Document()
    d.add_paragraph("First paragraph.")
    d.add_paragraph("Second paragraph.")
    d.save(str(f))
    text = extract(f).text
    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_corrupt_pdf_raises(tmp_path):
    f = tmp_path / "bad.pdf"
    f.write_bytes(b"this is not a pdf at all")
    with pytest.raises(ExtractionError):
        extract(f)


def test_unsupported_extension(tmp_path):
    f = tmp_path / "image.png"
    f.write_bytes(b"\x89PNG")
    with pytest.raises(ExtractionError):
        extract(f)
